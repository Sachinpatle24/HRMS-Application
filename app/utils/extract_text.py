# app/utils/extract_text.py

import fitz
import io  
import subprocess
import shutil
from docx import Document
import os
import tempfile
from pathlib import Path
from app.core.logger import get_custom_logger

logger = get_custom_logger(app_name="extract_text")

def clean_resume_text(text: str) -> str:
    if not text:
        return ""
    text = text.encode('ascii', 'ignore').decode('ascii')
    return text


def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from PDF using PyMuPDF."""
    try:
        pdf = fitz.open(file_path)
        text = "\n".join([page.get_text() for page in pdf])
        if not text.strip():
            logger.warning(f"PDF has no extractable text (likely scanned/image-only): {Path(file_path).name}")
            return ""
        return text
    except Exception as e:
        logger.exception(f"PDF extraction failed for {file_path}: {e}")
        return ""


def extract_text_from_docx(file_path: str) -> str:
    """Extract text from DOCX including tables, headers, and footers."""
    try:
        doc = Document(file_path)
        text_parts = []
        
        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)
        
        # Extract headers
        for section in doc.sections:
            header = section.header
            for para in header.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Extract footer
            footer = section.footer
            for para in footer.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
        
        text = "\n".join(text_parts)
        if not text.strip():
            logger.warning(f"DOCX has no extractable text: {Path(file_path).name}")
            return ""
        return text
    except Exception as e:
        logger.exception(f"DOCX extraction failed for {file_path}: {e}")
        return ""


def extract_text_from_doc(file_path: str) -> str:
    """
    Extract text from legacy .doc using LibreOffice headless conversion.
    Returns empty string if conversion fails or no text is found.
    """
    tmp_dir = None
    soffice_path = shutil.which("soffice")
    if not soffice_path:
        logger.error("LibreOffice 'soffice' not found in PATH. DOC extraction will fail.")
        return ""
    
    try:
        logger.info(f"Extracting text from DOC via LibreOffice: {file_path}")
        tmp_dir = tempfile.mkdtemp(prefix="doc_extract_")

        # Run LibreOffice conversion
        result = subprocess.run(
            [soffice_path, "--headless", "--convert-to", "txt:Text", "--outdir", str(tmp_dir), file_path],
            capture_output=True,
            text=True,
            timeout=30
        )

        if result.returncode != 0:
            logger.error(f"LibreOffice conversion failed: {result.stderr.strip()}")
            return ""

        txt_files = list(Path(tmp_dir).glob("*.txt"))
        if not txt_files:
            # logger.warning(f"No .txt output file created in {tmp_dir}")
            logger.warning(
                "LibreOffice produced no txt output",
                extra={"stderr": result.stderr, "stdout": result.stdout}
            )
            return ""

        txt_path = max(txt_files, key=lambda p: p.stat().st_size)
        text = txt_path.read_text(encoding="utf-8", errors="ignore")
        if not text.strip():
            logger.warning(f"DOC converted but empty: {Path(file_path).name}")
            return ""

        return text

    except subprocess.TimeoutExpired:
        logger.error(f"LibreOffice timeout for {file_path}")
        return ""
    except Exception as e:
        logger.exception(f"DOC extraction failed for {file_path}: {e}")
        return ""
    finally:
        # Cleanup temp files
        if tmp_dir:
            for f in Path(tmp_dir).glob("*"):
                try:
                    f.unlink()
                except Exception:
                    pass
            try:
                os.rmdir(tmp_dir)
            except Exception:
                pass


def extract_text_from_doc_bytes(file_data: bytes) -> str:
    """
    Extract text from legacy .doc bytes using LibreOffice headless conversion.
    Writes bytes to temporary .doc file and reuses extract_text_from_doc.
    """
    tmp_file = None
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".doc") as tmp:
            tmp.write(file_data)
            tmp_file = tmp.name

        return extract_text_from_doc(tmp_file)

    finally:
        if tmp_file and os.path.exists(tmp_file):
            os.remove(tmp_file)


def extract_text_from_file(file_path: str) -> str:
    """
    Unified extractor for PDF and DOCX.
    Returns plain text string.
    """
    try:
        ext = Path(file_path).suffix.lower()

        if ext == ".pdf":
            logger.info(f"Extracting text from PDF: {file_path}")
            return extract_text_from_pdf(file_path)

        elif ext == ".docx":
            logger.info(f"Extracting text from DOCX: {file_path}")
            return extract_text_from_docx(file_path)

        elif ext == ".doc":
            logger.info(f"Extracting text from DOC: {file_path}")
            return extract_text_from_doc(file_path)
        
        else:
            logger.warning(f"Unsupported file type: {ext} in {file_path}")
            return None

    except Exception as e:
        logger.exception(f"Unexpected error extracting file {file_path}: {e}")
        return ""

def extract_text_from_bytes(file_data: bytes, file_type: str) -> str:
    if file_type == "pdf":
        doc = fitz.open(stream=file_data, filetype="pdf")
        text = "\n".join(page.get_text() for page in doc)
        
        if not text.strip():
            logger.warning("PDF appears image-only/scanned")
            return ""
        
        return text

    if file_type == "docx":
        doc = Document(io.BytesIO(file_data))
        text_parts = []
    
        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)
        
        # Extract headers/footers
        for section in doc.sections:
            for para in section.header.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            for para in section.footer.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
        
        text = "\n".join(text_parts)
        if not text.strip():
            logger.warning("DOCX has no text content")
            return "" 
        
        return text
    
    if file_type == "doc":
        return extract_text_from_doc_bytes(file_data)

    raise ValueError(f"Unsupported file type: {file_type}")
